"""
Document Parser Service
Handles extraction of text from various document formats (PDF, DOCX, TXT)
"""

import io
import logging
from typing import Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service for parsing and extracting text from documents"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx', '.doc'}
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available"""
        self.pdf_available = False
        self.docx_available = False
        self.ocr_available = False
        
        try:
            import pdfplumber
            self.pdf_available = True
        except ImportError:
            logger.warning("pdfplumber not installed. PDF support disabled.")
        
        try:
            from docx import Document
            self.docx_available = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX support disabled.")
        
        try:
            import pytesseract
            from PIL import Image
            self.ocr_available = True
        except ImportError:
            logger.warning("pytesseract/Pillow not installed. OCR support disabled.")
    
    async def parse(self, file_content: bytes, filename: str) -> Tuple[str, dict]:
        """
        Parse document and extract text
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        extension = Path(filename).suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        metadata = {
            "filename": filename,
            "extension": extension,
            "size_bytes": len(file_content),
            "pages": None,
            "extraction_method": None
        }
        
        if extension == '.pdf':
            text, meta = await self._parse_pdf(file_content)
        elif extension == '.txt':
            text, meta = await self._parse_text(file_content)
        elif extension in ['.docx', '.doc']:
            text, meta = await self._parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        metadata.update(meta)
        return text, metadata
    
    async def _parse_pdf(self, content: bytes) -> Tuple[str, dict]:
        """Extract text from PDF file"""
        if not self.pdf_available:
            raise RuntimeError("PDF parsing not available. Install pdfplumber.")
        
        import pdfplumber
        
        text_parts = []
        metadata = {"extraction_method": "pdfplumber"}
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            metadata["pages"] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                else:
                    # Try OCR if text extraction fails
                    if self.ocr_available:
                        ocr_text = await self._ocr_page(page)
                        if ocr_text:
                            text_parts.append(f"--- Page {page_num} (OCR) ---\n{ocr_text}")
                            metadata["extraction_method"] = "pdfplumber+ocr"
        
        return "\n\n".join(text_parts), metadata
    
    async def _parse_text(self, content: bytes) -> Tuple[str, dict]:
        """Extract text from plain text file"""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                return text, {"extraction_method": "text", "encoding": encoding}
            except UnicodeDecodeError:
                continue
        
        # Fallback with error replacement
        text = content.decode('utf-8', errors='replace')
        return text, {"extraction_method": "text", "encoding": "utf-8-fallback"}
    
    async def _parse_docx(self, content: bytes) -> Tuple[str, dict]:
        """Extract text from DOCX file"""
        if not self.docx_available:
            raise RuntimeError("DOCX parsing not available. Install python-docx.")
        
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts), {"extraction_method": "python-docx"}
    
    async def _ocr_page(self, page) -> Optional[str]:
        """Perform OCR on a PDF page"""
        if not self.ocr_available:
            return None
        
        try:
            import pytesseract
            from PIL import Image
            
            # Convert page to image
            image = page.to_image(resolution=300)
            pil_image = image.original
            
            # Perform OCR
            text = pytesseract.image_to_string(pil_image)
            return text if text.strip() else None
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return None
    
    def parse_text_directly(self, text: str) -> Tuple[str, dict]:
        """
        Process raw text input (for API endpoint that accepts text directly)
        
        Args:
            text: Raw text content
            
        Returns:
            Tuple of (text, metadata)
        """
        return text, {
            "filename": "direct_input.txt",
            "extension": ".txt",
            "size_bytes": len(text.encode('utf-8')),
            "extraction_method": "direct"
        }
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and normalizing
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        import re
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    @staticmethod
    def detect_language(text: str) -> str:
        """
        Simple language detection based on common words
        
        Args:
            text: Text to analyze
            
        Returns:
            ISO language code (e.g., 'en', 'pt', 'de')
        """
        # Common words in different languages
        language_markers = {
            'en': ['the', 'and', 'is', 'are', 'shall', 'must', 'tax', 'rate', 'income'],
            'pt': ['o', 'a', 'de', 'da', 'do', 'imposto', 'taxa', 'alíquota', 'renda'],
            'es': ['el', 'la', 'de', 'del', 'impuesto', 'tasa', 'renta', 'gravamen'],
            'de': ['der', 'die', 'das', 'und', 'steuer', 'satz', 'einkommen', 'betrag'],
            'fr': ['le', 'la', 'de', 'du', 'impôt', 'taux', 'revenu', 'taxe'],
            'it': ['il', 'la', 'di', 'del', 'imposta', 'tasso', 'reddito', 'aliquota'],
        }
        
        text_lower = text.lower()
        words = set(text_lower.split())
        
        scores = {}
        for lang, markers in language_markers.items():
            score = sum(1 for marker in markers if marker in words)
            scores[lang] = score
        
        # Default to English if no clear winner
        best_lang = max(scores, key=scores.get) if scores else 'en'
        return best_lang if scores.get(best_lang, 0) > 0 else 'en'
